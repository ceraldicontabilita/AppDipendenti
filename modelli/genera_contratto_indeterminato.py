"""
Genera un template CORRETTO del «Contratto di lavoro subordinato a tempo
indeterminato» per Ceraldi Group (settore bar/pasticceria, CCNL Pubblici
Esercizi/Turismo H05Y).

Riproduce fedelmente il testo del contratto aziendale esistente correggendo gli
errori riscontrati nel modello in uso:
  - «Sig.ra Vincenzo Ceraldi» -> «Sig. Vincenzo Ceraldi» (legale rapp. uomo);
  - rimosso il refuso «52 riposi infrasettimanali»;
  - periodo di prova reso parametrico ({{periodo_prova}}) anziché fisso «15 giorni»;
  - eliminato l'artefatto «Napoli Barista.» nella riga luogo/data;
  - aggiunti retribuzione mensile, mensilità e ferie come segnaposto;
  - tutti i dati del dipendente come segnaposto nominali {{chiave}}.

Uso:  python modelli/genera_contratto_indeterminato.py
Output: modelli/Contratto_indeterminato_Turismo.docx
Poi: Assunzione -> Modelli contratto -> «Contratto a Tempo Indeterminato» -> Sostituisci.

NB: testo da far validare al consulente del lavoro prima dell'uso.
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "Contratto_indeterminato_Turismo.docx")


def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("CONTRATTO DI LAVORO SUBORDINATO A TEMPO INDETERMINATO")
    tr.bold = True
    tr.font.size = Pt(14)

    body(doc, "Datore di lavoro: Ceraldi Group S.r.l. in persona del legale rappresentante "
              "Sig. Vincenzo Ceraldi, con sede legale in Napoli alla Piazza Nazionale, n. 46, "
              "C.F. e partita IVA 04523831214.")
    body(doc, "Lavoratore: {{nome_completo}}, nato a {{luogo_nascita}} il {{data_nascita}}, "
              "residente in {{indirizzo}} con codice fiscale {{codice_fiscale}}.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CONVENGONO").bold = True

    h(doc, "1. OGGETTO, MANSIONI, INQUADRAMENTO")
    body(doc, "Il Sig. {{nome_completo}} è assunto da Ceraldi Group S.r.l. con assegnazione delle "
              "seguenti mansioni: {{mansione}}, inquadrato nel livello {{livello}} e con qualifica "
              "{{qualifica}} del CCNL Pubblici Esercizi, Ristorazione Collettiva e Commerciale e "
              "Turismo (Confcommercio-FIPE, codice CNEL H05Y). L'assunzione è a tempo indeterminato "
              "e decorre dal {{data_inizio}}.")

    h(doc, "2. PATTO DI PROVA")
    body(doc, "L'assunzione è subordinata al positivo superamento di un periodo di prova di "
              "{{periodo_prova}} giorni di calendario, durante il quale ciascuna delle parti sarà "
              "libera di recedere dal contratto senza obbligo di preavviso, nei limiti e termini "
              "previsti dal CCNL applicato.")

    h(doc, "3. LUOGO E ORARIO DI LAVORO")
    body(doc, "Il Lavoratore è assegnato ai locali del Datore di lavoro siti in Piazza Carità, 14 — "
              "80134 Napoli. Resta inteso che l'attività lavorativa potrà essere svolta "
              "temporaneamente anche in luoghi diversi da quello di assunzione (ad es. catering), "
              "come previsto dal CCNL di riferimento.")
    body(doc, "L'orario di lavoro è di {{ore_settimanali}} ore settimanali, con riposo settimanale "
              "di norma domenicale, secondo il CCNL applicato. L'orario aziendale potrà essere "
              "modificato per esigenze organizzative, previa comunicazione al Lavoratore nei "
              "termini previsti dal CCNL di riferimento.")

    h(doc, "4. TRATTAMENTO ECONOMICO")
    body(doc, "Il trattamento economico è pari ad euro {{stipendio_orario}} l'ora (stabilito in "
              "ragione della qualifica e della categoria di appartenenza dal CCNL richiamato al "
              "punto 1), per una retribuzione lorda mensile indicativa di euro {{stipendio_mensile}}. "
              "Sono corrisposte le seguenti mensilità aggiuntive: {{mensilita}}.")
    body(doc, "Per quanto concerne il lavoro straordinario, l'azienda inserisce in busta paga le ore "
              "extra effettuate e ne rilascia ricevuta a comprova.")

    h(doc, "5. FERIE E PERMESSI")
    body(doc, "Al Lavoratore spettano {{ferie_giorni}} giorni di ferie annue, oltre ai permessi "
              "retribuiti previsti dal CCNL. Buono pasto: {{ticket}}")

    h(doc, "6. REGOLAMENTO DISCIPLINARE")
    body(doc, "Il Lavoratore dichiara di essere a conoscenza delle norme relative alle infrazioni "
              "disciplinari, alle procedure di contestazione e alle sanzioni contenute nel Codice "
              "civile, nella l. n. 300/1970 e nel CCNL richiamato al punto 1, del quale dichiara di "
              "prendere visione in estratto, unitamente al regolamento aziendale in allegato. Il "
              "Lavoratore si impegna ad attenersi al regolamento aziendale e alle disposizioni "
              "interne adottate in Azienda.")

    h(doc, "7. SICUREZZA")
    body(doc, "Il Datore di lavoro dichiara di applicare le norme in materia di sicurezza sul lavoro, "
              "in particolare il d.lgs. n. 81/2008 e s.m.i. Il Lavoratore si impegna a uniformarsi "
              "alle relative prescrizioni e a segnalare eventuali situazioni anomale.")

    h(doc, "8. REGISTRAZIONI OBBLIGATORIE")
    body(doc, "Con l'assunzione il Lavoratore è iscritto nel Libro Unico del Lavoro tenuto ai sensi "
              "di legge.")

    h(doc, "9. PRIVACY")
    body(doc, "I dati del Lavoratore e, se del caso, dei suoi familiari, sono trattati ai sensi della "
              "normativa vigente ai soli fini della gestione del rapporto di lavoro, inclusi i "
              "rapporti con enti previdenziali, assistenziali e con l'amministrazione finanziaria.")

    h(doc, "10. RISERVATEZZA E OBBLIGO DI FEDELTÀ")
    body(doc, "Il Lavoratore si impegna alla massima riservatezza su dati e notizie di cui venga a "
              "conoscenza in occasione dell'attività lavorativa e si obbliga a non trattare affari in "
              "concorrenza con il Datore di lavoro ai sensi dell'art. 2105 c.c.")

    h(doc, "11. PREVIDENZA COMPLEMENTARE E TFR")
    body(doc, "Ai fini della destinazione del TFR si allega l'informativa ex art. 8, comma 8, del "
              "d.lgs. n. 252/2005 con la relativa modulistica. Il Lavoratore potrà destinare il TFR a "
              "un fondo di previdenza complementare (es. Fon.Te.) ovvero lasciarne la maturazione in "
              "azienda.")

    h(doc, "12. CLAUSOLA FINALE")
    body(doc, "Per quanto qui non espressamente previsto, il rapporto è regolato dal CCNL applicato e "
              "richiamato al punto 1 e dalle norme di legge in materia di lavoro e previdenza.")

    doc.add_paragraph()
    body(doc, "Napoli, lì ______________________")
    doc.add_paragraph()
    body(doc, "Il Datore di lavoro\t\t\t\t\tIl Lavoratore")
    body(doc, "____________________________\t\t\t____________________________")

    doc.add_paragraph()
    h(doc, "ALLEGATI")
    body(doc, "• Informativa sintetica sulle condizioni applicabili al rapporto di lavoro (D.Lgs. 152/1997).")
    body(doc, "• Informativa privacy ai sensi del Reg. UE 2016/679.")
    body(doc, "• Copia del regolamento disciplinare aziendale.")

    doc.save(OUT)
    print("Generato:", OUT)


if __name__ == "__main__":
    build()
