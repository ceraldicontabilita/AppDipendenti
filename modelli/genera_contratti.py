"""
Genera i template CORRETTI dei 4 contratti Ceraldi (settore bar/pasticceria,
CCNL Pubblici Esercizi/Turismo H05Y), in un unico posto (niente doppioni):

  - Contratto_indeterminato_Turismo.docx
  - Contratto_determinato_Turismo.docx
  - Contratto_part_time_indeterminato_Turismo.docx
  - Contratto_part_time_determinato_Turismo.docx

Riproduce fedelmente il testo aziendale correggendo gli errori del modello in uso
(«Sig.ra»→«Sig.», refuso «52 riposi», artefatto «Napoli Barista», periodo di
prova fisso) e usando segnaposto nominali {{chiave}} per tutti i dati variabili.

Uso:  python modelli/genera_contratti.py
Poi: Assunzione -> Modelli contratto -> Sostituisci (per ciascun tipo).
NB: testo da far validare al consulente del lavoro prima dell'uso.
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(__file__)

# tipo -> (titolo, decorrenza, clausola orario, file)
TIPI = {
    "indeterminato": {
        "titolo": "CONTRATTO DI LAVORO SUBORDINATO A TEMPO INDETERMINATO",
        "decorrenza": "L'assunzione è a tempo indeterminato e decorre dal {{data_inizio}}.",
        "orario": "L'orario di lavoro è di {{ore_settimanali}} ore settimanali, con riposo "
                  "settimanale di norma domenicale, secondo il CCNL applicato.",
        "file": "Contratto_indeterminato_Turismo.docx",
    },
    "determinato": {
        "titolo": "CONTRATTO DI LAVORO SUBORDINATO A TEMPO DETERMINATO",
        "decorrenza": "L'assunzione è a tempo determinato e decorre dal {{data_inizio}} al "
                      "{{data_fine}}, per le ragioni indicate dalle parti nel rispetto del "
                      "D.Lgs. n. 81/2015 e del CCNL applicato.",
        "orario": "L'orario di lavoro è di {{ore_settimanali}} ore settimanali, con riposo "
                  "settimanale di norma domenicale, secondo il CCNL applicato.",
        "file": "Contratto_determinato_Turismo.docx",
    },
    "part_time_ind": {
        "titolo": "CONTRATTO DI LAVORO SUBORDINATO PART-TIME A TEMPO INDETERMINATO",
        "decorrenza": "L'assunzione è a tempo indeterminato e a tempo parziale e decorre dal "
                      "{{data_inizio}}.",
        "orario": "L'orario di lavoro è di {{ore_settimanali}} ore settimanali a tempo parziale, "
                  "secondo la distribuzione concordata tra le parti, nel rispetto del CCNL applicato.",
        "file": "Contratto_part_time_indeterminato_Turismo.docx",
    },
    "part_time_det": {
        "titolo": "CONTRATTO DI LAVORO SUBORDINATO PART-TIME A TEMPO DETERMINATO",
        "decorrenza": "L'assunzione è a tempo determinato e a tempo parziale e decorre dal "
                      "{{data_inizio}} al {{data_fine}}, nel rispetto del D.Lgs. n. 81/2015 e del "
                      "CCNL applicato.",
        "orario": "L'orario di lavoro è di {{ore_settimanali}} ore settimanali a tempo parziale, "
                  "secondo la distribuzione concordata tra le parti, nel rispetto del CCNL applicato.",
        "file": "Contratto_part_time_determinato_Turismo.docx",
    },
}


def _h(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; return p


def _b(doc, text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(6); return p


def build(tipo: str):
    cfg = TIPI[tipo]
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(cfg["titolo"]); tr.bold = True; tr.font.size = Pt(14)

    _b(doc, "Datore di lavoro: Ceraldi Group S.r.l. in persona del legale rappresentante "
            "Sig. Vincenzo Ceraldi, con sede legale in Napoli alla Piazza Nazionale, n. 46, "
            "C.F. e partita IVA 04523831214.")
    _b(doc, "Lavoratore: {{nome_completo}}, nato a {{luogo_nascita}} il {{data_nascita}}, "
            "residente in {{indirizzo}} con codice fiscale {{codice_fiscale}}.")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CONVENGONO").bold = True

    _h(doc, "1. OGGETTO, MANSIONI, INQUADRAMENTO")
    _b(doc, "Il Sig. {{nome_completo}} è assunto da Ceraldi Group S.r.l. con assegnazione delle "
            "seguenti mansioni: {{mansione}}, inquadrato nel livello {{livello}} e con qualifica "
            "{{qualifica}} del CCNL Pubblici Esercizi, Ristorazione Collettiva e Commerciale e "
            "Turismo (Confcommercio-FIPE, codice CNEL H05Y). " + cfg["decorrenza"])

    _h(doc, "2. PATTO DI PROVA")
    _b(doc, "L'assunzione è subordinata al positivo superamento di un periodo di prova di "
            "{{periodo_prova}} giorni di calendario, durante il quale ciascuna delle parti sarà "
            "libera di recedere senza obbligo di preavviso, nei limiti e termini previsti dal CCNL "
            "applicato (per i contratti a termine non oltre il 50% della durata del rapporto).")

    _h(doc, "3. LUOGO E ORARIO DI LAVORO")
    _b(doc, "Il Lavoratore è assegnato ai locali del Datore di lavoro siti in Piazza Carità, 14 — "
            "80134 Napoli. L'attività potrà essere svolta temporaneamente anche in luoghi diversi "
            "(ad es. catering), come previsto dal CCNL di riferimento.")
    _b(doc, cfg["orario"] + " L'orario aziendale potrà essere modificato per esigenze "
            "organizzative, previa comunicazione al Lavoratore nei termini previsti dal CCNL.")

    _h(doc, "4. TRATTAMENTO ECONOMICO")
    _b(doc, "Il trattamento economico è pari ad euro {{stipendio_orario}} l'ora (stabilito in "
            "ragione della qualifica e della categoria di appartenenza dal CCNL richiamato al "
            "punto 1), per una retribuzione lorda mensile indicativa di euro {{stipendio_mensile}}. "
            "Sono corrisposte le seguenti mensilità aggiuntive: {{mensilita}}.")
    _b(doc, "Per il lavoro straordinario l'azienda inserisce in busta paga le ore extra effettuate "
            "e ne rilascia ricevuta a comprova.")

    _h(doc, "5. FERIE E PERMESSI")
    _b(doc, "Al Lavoratore spettano {{ferie_giorni}} giorni di ferie annue, oltre ai permessi "
            "retribuiti previsti dal CCNL. Buono pasto: {{ticket}}")

    _h(doc, "6. REGOLAMENTO DISCIPLINARE")
    _b(doc, "Il Lavoratore dichiara di conoscere le norme su infrazioni disciplinari, procedure di "
            "contestazione e sanzioni contenute nel Codice civile, nella l. n. 300/1970 e nel CCNL "
            "richiamato al punto 1, del quale dichiara di prendere visione in estratto unitamente al "
            "regolamento aziendale in allegato, impegnandosi a osservarlo.")

    _h(doc, "7. SICUREZZA")
    _b(doc, "Il Datore di lavoro applica le norme in materia di sicurezza sul lavoro, in particolare "
            "il d.lgs. n. 81/2008 e s.m.i. Il Lavoratore si impegna a uniformarsi alle prescrizioni e "
            "a segnalare situazioni anomale.")

    _h(doc, "8. REGISTRAZIONI OBBLIGATORIE E PRIVACY")
    _b(doc, "Con l'assunzione il Lavoratore è iscritto nel Libro Unico del Lavoro. I dati personali "
            "sono trattati ai sensi del Reg. UE 2016/679 ai soli fini della gestione del rapporto di "
            "lavoro, come da informativa privacy allegata.")

    _h(doc, "9. RISERVATEZZA E OBBLIGO DI FEDELTÀ")
    _b(doc, "Il Lavoratore si impegna alla riservatezza su dati e notizie acquisiti in occasione "
            "dell'attività lavorativa e a non trattare affari in concorrenza con il Datore di lavoro "
            "ai sensi dell'art. 2105 c.c.")

    _h(doc, "10. PREVIDENZA COMPLEMENTARE E TFR")
    _b(doc, "Ai fini della destinazione del TFR si allega l'informativa ex art. 8, comma 8, del "
            "d.lgs. n. 252/2005. Il Lavoratore potrà destinare il TFR a un fondo di previdenza "
            "complementare (es. Fon.Te.) o lasciarne la maturazione in azienda.")

    _h(doc, "11. CLAUSOLA FINALE")
    _b(doc, "Per quanto qui non previsto, il rapporto è regolato dal CCNL applicato e richiamato al "
            "punto 1 e dalle norme di legge in materia di lavoro e previdenza.")

    doc.add_paragraph()
    _b(doc, "Napoli, lì ______________________")
    doc.add_paragraph()
    _b(doc, "Il Datore di lavoro\t\t\t\t\tIl Lavoratore")
    _b(doc, "____________________________\t\t\t____________________________")

    doc.add_paragraph()
    _h(doc, "ALLEGATI")
    _b(doc, "• Informativa sulle condizioni applicabili al rapporto di lavoro (D.Lgs. 152/1997).")
    _b(doc, "• Informativa privacy (Reg. UE 2016/679).")
    _b(doc, "• Regolamento interno/disciplinare aziendale.")

    out = os.path.join(HERE, cfg["file"])
    doc.save(out)
    return out


if __name__ == "__main__":
    for tipo in TIPI:
        print("Generato:", build(tipo))
