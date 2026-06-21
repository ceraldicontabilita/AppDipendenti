"""
Genera l'INFORMATIVA AI SENSI DEL D.LGS. 152/1997 adattata al CCNL
"Pubblici Esercizi, Ristorazione Collettiva e Commerciale e Turismo"
(Confcommercio-FIPE), codice CNEL H05Y — settore bar/pasticceria.

Sostituisce l'informativa Terziario precedente. Mappa applicata:
  - CCNL Terziario  -> CCNL Pubblici Esercizi/Turismo (H05Y, rinnovo 5/6/2024)
  - Ebiter/EBT      -> EBNT + Ente Bilaterale Territoriale Campania
  - Fondo EST       -> resta (assistenza sanitaria, obbligatorio anche per H05Y)
  - Fon.Te.         -> resta (previdenza complementare, copre il Turismo)
  - Rimosse le voci solo-Terziario (Terzo elemento, indennità liv. Q,
    superminimo VII liv., rimandi artt. 241-248).

I dati del dipendente sono segnaposto nominali {{chiave}} compilati dal modulo
Assunzione (fill_contract_template). Gli IMPORTI TABELLARI non sono inventati:
restano da compilare/validare col consulente del lavoro sulle tabelle ufficiali
del CCNL H05Y per il livello di inquadramento.

Uso:  python modelli/genera_informativa_152.py
Output: modelli/INFORMATIVA_152_1997_Turismo.docx
Poi: caricalo nel sistema come template "Informativa D.Lgs. 152/1997"
(Assunzione -> Modelli contratto -> Sostituisci).
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "INFORMATIVA_152_1997_Turismo.docx")

# Segnaposto che richiedono la validazione del consulente del lavoro.
TODO = RGBColor(0xC4, 0x89, 0x4A)  # sand/warn — evidenzia le voci da confermare


def h(doc, text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def todo(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = TODO
    return p


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("INFORMATIVA AI SENSI DEL D.LGS. 152/1997")
    tr.bold = True
    tr.font.size = Pt(15)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("(condizioni applicabili al rapporto di lavoro — obblighi di trasparenza, D.Lgs. 104/2022)")
    sr.italic = True

    body(doc, "Ceraldi Group S.r.l., con sede in Napoli, in qualità di datore di lavoro, "
              "fornisce al lavoratore le seguenti informazioni sul rapporto di lavoro.")

    h(doc, "1. Identità delle parti")
    body(doc, "Datore di lavoro: Ceraldi Group S.r.l. — Napoli.")
    body(doc, "Lavoratore: {{nome_completo}}, nato a {{luogo_nascita}} il {{data_nascita}}, "
              "residente in {{indirizzo}}, codice fiscale {{codice_fiscale}}.")

    h(doc, "2. Luogo di lavoro")
    body(doc, "La prestazione si svolge presso le unità operative dell'azienda (attività di "
              "bar/pasticceria). In mancanza di sede fissa o prevalente, il lavoro si intende "
              "svolto in luoghi diversi secondo le esigenze aziendali.")

    h(doc, "3. Data di inizio e durata del rapporto")
    body(doc, "Inizio del rapporto: {{data_inizio}}. Per i rapporti a tempo determinato la "
              "scadenza è fissata al {{data_fine}}.")

    h(doc, "4. Contratto collettivo applicato e inquadramento")
    body(doc, "Al rapporto si applica il CCNL «Pubblici Esercizi, Ristorazione Collettiva e "
              "Commerciale e Turismo» (Confcommercio-FIPE), codice CNEL H05Y, rinnovo del "
              "5/6/2024, con scadenza 31/12/2027, e successive modifiche.")
    body(doc, "Inquadramento (art. 101 CCNL — classificazione del personale): livello {{livello}}, "
              "qualifica {{qualifica}}, mansione {{mansione}}.")

    h(doc, "5. Periodo di prova")
    body(doc, "Il periodo di prova è di {{periodo_prova}} giorni, secondo quanto previsto dal CCNL "
              "in funzione del livello di inquadramento (per i contratti a termine non oltre il 50% "
              "della durata del rapporto).")

    h(doc, "6. Retribuzione")
    body(doc, "Paga oraria: euro {{stipendio_orario}} l'ora, per {{ore_settimanali}} ore settimanali. "
              "Retribuzione lorda mensile indicativa: euro {{stipendio_mensile}}.")
    body(doc, "Mensilità: il CCNL prevede 14 mensilità — 13ª corrisposta a dicembre e 14ª a luglio "
              "(esclusi assegni familiari e scatti).")
    todo(doc, "[Da compilare/validare col consulente del lavoro: minimi tabellari del CCNL H05Y "
              "per il livello {{livello}} — paga base, contingenza, eventuali scatti di anzianità. "
              "Non riportare importi non verificati sulle tabelle ufficiali vigenti.]")

    h(doc, "7. Orario di lavoro")
    body(doc, "L'orario normale è di 40 ore settimanali, distribuite di norma su 5 giornate e mezza, "
              "secondo l'organizzazione aziendale e gli istituti previsti dal CCNL.")

    h(doc, "8. Ferie e permessi")
    body(doc, "Ferie: {{ferie_giorni}} giorni lavorativi di ferie all'anno. Permessi retribuiti: "
              "dopo 4 anni di servizio 104 ore annue (comprensive di 32 ore per ex festività), "
              "secondo il CCNL.")

    h(doc, "9. Termini di preavviso")
    todo(doc, "[Da indicare secondo le tabelle del CCNL H05Y in base a livello e anzianità: "
              "i termini di preavviso per recesso. Valore da confermare col consulente.]")

    h(doc, "10. Enti bilaterali e fondi contrattuali")
    body(doc, "Ente bilaterale: EBNT — Ente Bilaterale Nazionale Turismo, unitamente all'Ente "
              "Bilaterale Territoriale della Campania.")
    body(doc, "Assistenza sanitaria integrativa: Fondo EST (iscrizione obbligatoria prevista dal CCNL).")
    body(doc, "Previdenza complementare: Fondo Fon.Te., a cui il lavoratore può aderire secondo le "
              "modalità contrattuali.")

    h(doc, "11. Trattamento dei dati personali")
    body(doc, "Il trattamento dei dati personali del lavoratore avviene nel rispetto del Reg. UE "
              "2016/679 (GDPR) e del D.Lgs. 196/2003, come da specifica informativa privacy consegnata "
              "separatamente.")

    doc.add_paragraph()
    body(doc, "Luogo e data: ______________________")
    doc.add_paragraph()
    body(doc, "Il datore di lavoro — Ceraldi Group S.r.l.        Per presa visione, il lavoratore")
    body(doc, "____________________________                      ____________________________")

    doc.save(OUT)
    print("Generato:", OUT)


if __name__ == "__main__":
    build()
