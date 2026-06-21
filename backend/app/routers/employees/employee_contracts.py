"""
Employee Contracts Router - Gestione contratti dipendenti.
"""
from fastapi import APIRouter, HTTPException, Body, UploadFile, File, Depends
from fastapi.responses import FileResponse, Response, StreamingResponse
from typing import Dict, Any, List, Optional
from datetime import datetime, timezone
import logging
import os
import io
import ssl
import smtplib
import base64
import uuid
import shutil
from email.message import EmailMessage
from docx import Document
import tempfile

from backend.app.database import Database, Collections
from backend.app.utils.error_handler import handle_errors
from backend.app.services.openapi_signature import (
    get_client, OpenAPIConfigError, OpenAPIError,
)

logger = logging.getLogger(__name__)
router = APIRouter()

COLL_TEMPLATES = "contract_templates"   # template .docx persistenti (MongoDB-first)

# Directory effimere (solo file temporanei di lavorazione): /tmp è scrivibile su Render.
CONTRACTS_DIR = "/tmp/uploads/contracts"
TEMPLATES_DIR = "/tmp/uploads/contract_templates"

# Available contract types
CONTRACT_TYPES = [
    {"id": "determinato", "name": "Contratto a Tempo Determinato", "filename": "Contratto derminato.docx"},
    {"id": "indeterminato", "name": "Contratto a Tempo Indeterminato", "filename": "Contratto indetermionato.docx"},
    {"id": "part_time_det", "name": "Contratto Part-Time Determinato", "filename": "Contratto part_time determinato.docx"},
    {"id": "part_time_ind", "name": "Contratto Part-Time Indeterminato", "filename": "Contratto part_time indeterminato.docx"},
    {"id": "informativa_152", "name": "Informativa D.Lgs. 152/1997", "filename": "INFORMATIVA AI SENSI DEL D.LGS. 152-1997.docx"},
    {"id": "informativa_privacy", "name": "Informativa Privacy", "filename": "Informativa-Privacy.docx"},
    {"id": "regolamento", "name": "Regolamento Interno Aziendale", "filename": "REGOLAMENTO INTERNO AZIENDALE.docx"},
    {"id": "richiesta_ferie", "name": "Richiesta Ferie", "filename": "RICHIESTA FERIE.docx"},
]


def ensure_dirs():
    """Create directories if they don't exist (best-effort)."""
    for d in (CONTRACTS_DIR, TEMPLATES_DIR):
        try:
            os.makedirs(d, exist_ok=True)
        except Exception:
            pass


async def _resolve_template(ct: Dict[str, str]) -> str:
    """Risolve il percorso del template .docx.

    Priorità MongoDB-first: se il template è salvato in `contract_templates`
    (persistente tra i deploy), lo scrive in un file temporaneo e ritorna quel
    path; altrimenti usa il file su disco in TEMPLATES_DIR (effimero su Render).
    """
    ensure_dirs()
    db = Database.get_db()
    doc = await db[COLL_TEMPLATES].find_one({"tipo": ct["id"]}, {"_id": 0, "file_data": 1})
    if doc and doc.get("file_data"):
        tmp = tempfile.mktemp(suffix=".docx")
        with open(tmp, "wb") as f:
            f.write(base64.b64decode(doc["file_data"]))
        return tmp
    disk = os.path.join(TEMPLATES_DIR, ct["filename"])
    if os.path.exists(disk):
        return disk
    raise HTTPException(404, f"Template non caricato: {ct['name']}. Caricalo dalla sezione Assunzione.")


def _to_float(val: Any) -> Optional[float]:
    """Converte in float un valore numerico tollerando la virgola decimale IT."""
    if val is None or val == "":
        return None
    try:
        return float(str(val).replace("€", "").replace(",", ".").strip())
    except (ValueError, TypeError):
        return None


def _fmt_euro(val: Optional[float]) -> str:
    """Formatta un importo in stile italiano (1.234,56) senza simbolo."""
    if val is None:
        return "______"
    return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def compute_stipendio_mensile(stipendio_orario: Any, ore_settimanali: Any) -> Optional[float]:
    """Calcola il lordo mensile teorico: paga oraria × ore settimanali × 52 / 12."""
    orario = _to_float(stipendio_orario)
    ore = _to_float(ore_settimanali)
    if orario is None or ore is None:
        return None
    return round(orario * ore * 52 / 12, 2)


def fill_contract_template(template_path: str, employee_data: Dict[str, Any]) -> str:
    """
    Fill contract template with employee data.
    Replaces specific text patterns with employee data.

    Supporta due meccanismi di segnaposto, combinabili nello stesso .docx:
      1. Puntini di sospensione (… ……) compilati per posizione (legacy).
      2. Segnaposto nominali `{{chiave}}` (es. {{ore_settimanali}},
         {{stipendio_mensile}}, {{periodo_prova}}) — il modo consigliato per i
         nuovi campi CCNL Turismo, indipendente dal layout del documento.
    """
    doc = Document(template_path)

    # Build full name
    nome_completo = employee_data.get("nome_completo", "")
    if not nome_completo:
        nome_completo = f"{employee_data.get('cognome', '')} {employee_data.get('nome', '')}".strip()

    # Campi CCNL Pubblici Esercizi / Turismo (H05Y) — parametrici, non inventati.
    ore_settimanali = employee_data.get("ore_settimanali") or "40"
    stipendio_orario = employee_data.get("stipendio_orario") or employee_data.get("salary")
    mensile = compute_stipendio_mensile(stipendio_orario, ore_settimanali)
    ferie_giorni = employee_data.get("ferie_giorni") or "26"
    periodo_prova = employee_data.get("periodo_prova") or ""
    ticket_attivo = bool(employee_data.get("ticket_buono"))
    ticket_importo = employee_data.get("ticket_importo")
    tredicesima = employee_data.get("tredicesima", True)
    quattordicesima = employee_data.get("quattordicesima", True)

    # Frasi pronte da inserire nel .docx tramite segnaposto nominali.
    if ticket_attivo:
        _imp = _to_float(ticket_importo)
        ticket_txt = (
            f"Buono pasto di euro {_fmt_euro(_imp)} giornalieri, riconosciuto dopo 1 anno di servizio."
            if _imp is not None else
            "Buono pasto giornaliero riconosciuto dopo 1 anno di servizio."
        )
    else:
        ticket_txt = "Non previsto."
    mensilita_lista = []
    if tredicesima:
        mensilita_lista.append("13ª (corrisposta a dicembre)")
    if quattordicesima:
        mensilita_lista.append("14ª (corrisposta a luglio)")
    mensilita_txt = " e ".join(mensilita_lista) if mensilita_lista else "12 mensilità"

    # All values to replace
    def g(*keys, default="______"):
        """Primo valore non vuoto tra le chiavi date; evita di scrivere 'None'."""
        for k in keys:
            v = employee_data.get(k)
            if v not in (None, "", "None"):
                return v
        return default

    def _fmt_date(v, default="______"):
        """Formatta una data in gg/mm/aaaa; tollera ISO e valori già formattati."""
        if v in (None, "", "None"):
            return default
        s = str(v)
        try:
            return datetime.fromisoformat(s.replace("Z", "")).strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            return s

    data_values = {
        "nome_completo": nome_completo or "______",
        "cognome": g("cognome"),
        "nome": g("nome"),
        "codice_fiscale": g("codice_fiscale", "cf"),
        "data_nascita": _fmt_date(g("data_nascita", default=None)),
        "luogo_nascita": g("luogo_nascita", "comune_nascita", "citta_nascita"),
        "indirizzo": g("indirizzo", "residenza"),
        "mansione": g("mansione", "qualifica"),
        "livello": g("livello"),
        "qualifica": g("qualifica", "mansione"),
        "stipendio_orario": str(stipendio_orario) if stipendio_orario not in (None, "") else "______",
        "data_inizio": _fmt_date(g("data_inizio", "hire_date", default=None)),
        "data_fine": _fmt_date(g("data_fine", default=None), default=""),
        # Nuovi campi CCNL Turismo (usabili come {{chiave}} nei .docx)
        "ore_settimanali": str(ore_settimanali),
        "stipendio_mensile": _fmt_euro(mensile),
        "ferie_giorni": str(ferie_giorni),
        "periodo_prova": str(periodo_prova) if periodo_prova not in (None, "") else "______",
        "ticket": ticket_txt,
        "mensilita": mensilita_txt,
    }

    # Decorrenza: senza data fine (indeterminato) niente "al ..." in coda.
    _df = data_values["data_fine"]
    _decorr = f"decorre dal {data_values['data_inizio']}" + (f" al {_df}" if _df else "")

    # Alias accettati per i segnaposto nominali (tolleranza sui nomi nel .docx).
    named_aliases = {
        "ore": "ore_settimanali",
        "ore_lavoro": "ore_settimanali",
        "mensile": "stipendio_mensile",
        "stipendio_mese": "stipendio_mensile",
        "paga_mensile": "stipendio_mensile",
        "paga_oraria": "stipendio_orario",
        "ferie": "ferie_giorni",
        "prova": "periodo_prova",
        "buono_pasto": "ticket",
    }

    import re as _re

    def _apply_named(text: str) -> str:
        """Sostituisce i segnaposto nominali {{chiave}} (case-insensitive)."""
        if "{{" not in text:
            return text
        def _sub(m):
            key = m.group(1).strip().lower()
            key = named_aliases.get(key, key)
            return str(data_values.get(key, m.group(0)))
        return _re.sub(r"\{\{\s*([\w]+)\s*\}\}", _sub, text)

    def replace_placeholders(text: str) -> str:
        """Replace ellipsis placeholders with employee data."""
        result = _apply_named(text)
        # Periodo di prova parametrico anche su template con "15 giorni" fisso.
        pp = data_values["periodo_prova"]
        if pp and pp != "______":
            result = _re.sub(r'(prova di|minimo di)\s*\d+\s*giorni',
                             rf'\1 {pp} giorni', result, flags=_re.IGNORECASE)
        if "…" not in result:
            return result

        # The template uses Unicode ellipsis character (…) repeated multiple times
        # We need to replace these patterns specifically
        
        # Pattern 1: "Lavoratore: ……………, nato a …………. il ……………………, residente in ………………………………… con codice fiscale ……………………………."
        if "Lavoratore:" in result and "…" in result:
            # Replace the entire line (niente mansione davanti al nome)
            result = f"Lavoratore: {data_values['nome_completo']}, nato a {data_values['luogo_nascita']} il {data_values['data_nascita']}, residente in {data_values['indirizzo']} con codice fiscale {data_values['codice_fiscale']}."
        
        # Pattern 2: "IL Sig. ……………………………. è assunto" - this line contains EVERYTHING
        elif "IL Sig." in result and "è assunto" in result and "…" in result:
            import re
            # Replace name
            result = re.sub(r'IL Sig\.\s*[…\.]+\s*è assunto', f"IL Sig. {data_values['nome_completo']} è assunto", result)
            # Replace mansioni
            result = re.sub(r'mansioni:\s*[…\.]+\s*inquadrato', f"mansioni: {data_values['mansione']} inquadrato", result, flags=re.IGNORECASE)
            # Replace livello
            result = re.sub(r'livello\s*[…\.]+\s*e con', f"livello {data_values['livello']} e con", result, flags=re.IGNORECASE)
            # Replace qualifica
            result = re.sub(r'qualifica\s*[…\.]+\s*del', f"qualifica {data_values['qualifica']} del", result, flags=re.IGNORECASE)
            # Replace date decorrenza (senza "al" se indeterminato)
            result = re.sub(r'decorre dal\s*[…\.]+\s*al\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
            result = re.sub(r'decorre dal\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
        
        # Pattern 3: "mansioni: ………………………… inquadrato"
        elif "mansioni:" in result and "…" in result:
            import re
            result = re.sub(r'mansioni:\s*[…\.]+\s*inquadrato', f"mansioni: {data_values['mansione']} inquadrato", result)
        
        # Pattern 3b: "delle seguenti mansioni:" followed by placeholders
        elif "seguenti mansioni:" in result.lower() and "…" in result:
            import re
            result = re.sub(r'seguenti mansioni:\s*[…\.]+', f"seguenti mansioni: {data_values['mansione']}", result, flags=re.IGNORECASE)
        
        # Pattern 4: "livello …….." or "livello …………"
        elif "livello" in result.lower() and "…" in result:
            import re
            result = re.sub(r'livello\s*[…\.]+', f"livello {data_values['livello']}", result, flags=re.IGNORECASE)
        
        # Pattern 5: "qualifica …………" 
        elif "qualifica" in result.lower() and "…" in result:
            import re
            result = re.sub(r'qualifica\s*[…\.]+', f"qualifica {data_values['qualifica']}", result, flags=re.IGNORECASE)
        
        # Pattern 6: "euro ………………… ora" (stipendio)
        elif "euro" in result.lower() and "ora" in result.lower() and "…" in result:
            import re
            result = re.sub(r'euro\s*[…\.]+\s*ora', f"euro {data_values['stipendio_orario']} ora", result, flags=re.IGNORECASE)
        
        # Pattern 7: "decorre dal ………… al …………"
        elif "decorre dal" in result.lower() and "…" in result:
            import re
            result = re.sub(r'decorre dal\s*[…\.]+\s*al\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
            result = re.sub(r'decorre dal\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
        
        # Generic fallback: replace any remaining ellipsis sequences
        elif "…" in result:
            import re
            # Replace sequences of 10+ ellipsis with longer values
            result = re.sub(r'[…]{10,}', data_values['nome_completo'], result)
            # Replace sequences of 6-9 ellipsis with medium values
            result = re.sub(r'[…]{6,9}', data_values['mansione'], result)
            # Replace sequences of 3-5 ellipsis with shorter values
            result = re.sub(r'[…]{3,5}', "______", result)
        
        return result
    
    # Process all paragraphs (ogni paragrafo: i segnaposto e le sostituzioni
    # parametriche come il periodo di prova possono comparire anche senza puntini).
    for para in doc.paragraphs:
        new_text = replace_placeholders(para.text)
        if new_text != para.text:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run.text = ""
                first_run.text = new_text
            else:
                para.text = new_text

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text = replace_placeholders(para.text)
                    if new_text != para.text:
                        if para.runs:
                            first_run = para.runs[0]
                            for run in para.runs[1:]:
                                run.text = ""
                            first_run.text = new_text
                        else:
                            para.text = new_text
    
    # Save to temp file
    output_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_path)
    
    return output_path


@router.get("/types")
@handle_errors
async def get_contract_types() -> List[Dict[str, str]]:
    """Get available contract types."""
    return CONTRACT_TYPES


@router.get("/templates")
@handle_errors
async def list_templates() -> List[Dict[str, Any]]:
    """Elenco template con disponibilità (MongoDB-first, poi disco)."""
    ensure_dirs()
    db = Database.get_db()
    in_mongo = {d["tipo"] async for d in db[COLL_TEMPLATES].find({}, {"_id": 0, "tipo": 1})}
    templates = []
    for ct in CONTRACT_TYPES:
        exists = (ct["id"] in in_mongo) or os.path.exists(os.path.join(TEMPLATES_DIR, ct["filename"]))
        templates.append({
            "id": ct["id"],
            "name": ct["name"],
            "filename": ct["filename"],
            "available": exists,
        })
    return templates


@router.post("/template/{contract_type}")
@handle_errors
async def upload_template(contract_type: str, file: UploadFile = File(...)) -> Dict[str, Any]:
    """Carica/sostituisce un template .docx (salvato su MongoDB, persistente)."""
    ct = next((c for c in CONTRACT_TYPES if c["id"] == contract_type), None)
    if not ct:
        raise HTTPException(400, f"Tipo contratto non valido: {contract_type}")
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "File vuoto")
    if len(raw) > 12 * 1024 * 1024:
        raise HTTPException(400, "File troppo grande (max 12MB)")
    db = Database.get_db()
    await db[COLL_TEMPLATES].update_one(
        {"tipo": contract_type},
        {"$set": {
            "tipo": contract_type,
            "name": ct["name"],
            "filename": file.filename or ct["filename"],
            "file_data": base64.b64encode(raw).decode("utf-8"),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True)
    return {"ok": True, "tipo": contract_type, "name": ct["name"]}


@router.post("/generate/{employee_id}")
@handle_errors
async def generate_contract(employee_id: str, data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """
    Generate a contract for an employee.
    
    Request body:
    {
        "contract_type": "determinato",
        "additional_data": {
            "livello": "5",
            "stipendio_orario": "8.50",
            "qualifica": "Barista"
        }
    }
    """
    ensure_dirs()
    
    contract_type = data.get("contract_type") or data.get("contract_type_id")
    additional_data = data.get("additional_data", {})
    
    # Find contract type
    ct = next((c for c in CONTRACT_TYPES if c["id"] == contract_type), None)
    if not ct:
        raise HTTPException(status_code=400, detail=f"Tipo contratto non valido: {contract_type}")
    
    # Get employee
    db = Database.get_db()
    employee = await db[Collections.EMPLOYEES].find_one({"id": employee_id}, {"_id": 0})
    
    if not employee:
        raise HTTPException(status_code=404, detail="Dipendente non trovato")
    
    # Risolvi template (MongoDB-first, fallback su disco)
    template_path = await _resolve_template(ct)
    
    # Merge employee data with additional data
    employee_data = {**employee, **additional_data}
    
    # Format date if present
    if employee_data.get("data_nascita"):
        try:
            dt = datetime.fromisoformat(str(employee_data["data_nascita"]).replace("Z", ""))
            employee_data["data_nascita"] = dt.strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            pass
    
    try:
        # Generate filled contract
        output_path = fill_contract_template(template_path, employee_data)
        
        # Create final filename
        safe_name = employee_data.get("nome_completo", "dipendente").replace(" ", "_")
        final_filename = f"{ct['id']}_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        final_path = os.path.join(CONTRACTS_DIR, final_filename)
        
        # Move to contracts dir
        shutil.move(output_path, final_path)
        
        # Read file and encode to base64 for MongoDB (architettura MongoDB-first)
        import base64
        with open(final_path, 'rb') as f:
            file_content = f.read()
        file_base64 = base64.b64encode(file_content).decode('utf-8')
        
        # Lordo mensile teorico (paga oraria × ore settimanali × 52 / 12)
        mensile = compute_stipendio_mensile(
            employee_data.get("stipendio_orario") or employee_data.get("salary"),
            employee_data.get("ore_settimanali") or "40",
        )

        # Record contract generation with base64 content
        contract_record = {
            "id": str(uuid.uuid4()),
            "employee_id": employee_id,
            "employee_name": employee_data.get("nome_completo"),
            "contract_type": contract_type,
            "contract_name": ct["name"],
            "filename": final_filename,
            "filepath": final_path,
            "file_data": file_base64,  # Architettura MongoDB-first
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "stipendio_mensile": mensile,
            "additional_data": additional_data
        }

        await db["employee_contracts"].insert_one(contract_record.copy())

        return {
            "success": True,
            "message": f"Contratto generato per {employee_data.get('nome_completo')}",
            "stipendio_mensile": mensile,
            "contract": {
                "id": contract_record["id"],
                "filename": final_filename,
                "download_url": f"/api/contracts/download/{contract_record['id']}"
            }
        }
        
    except Exception as e:
        logger.error(f"Error generating contract: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Errore generazione contratto: {str(e)}")


@router.get("/download/{contract_id}")
@handle_errors
async def download_contract(contract_id: str):
    """
    Download a generated contract.
    Architettura MongoDB-first: priorità a file_data da MongoDB.
    """
    import base64
    from fastapi.responses import Response
    
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    # Priorità: file_data da MongoDB (architettura MongoDB-first)
    file_data = contract.get("file_data")
    if file_data:
        content = base64.b64decode(file_data)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{contract.get("filename", "contratto.docx")}"'}
        )
    
    # Fallback per contratti legacy con solo filepath
    filepath = contract.get("filepath")
    if filepath and os.path.exists(filepath):
        return FileResponse(
            filepath,
            filename=contract.get("filename"),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    raise HTTPException(status_code=404, detail="File contratto non trovato")


@router.get("/employee/{employee_id}")
@handle_errors
async def get_employee_contracts(employee_id: str) -> List[Dict[str, Any]]:
    """Get all contracts for an employee."""
    db = Database.get_db()
    contracts = await db["employee_contracts"].find(
        {"employee_id": employee_id},
        {"_id": 0}
    ).sort("generated_at", -1).to_list(100)
    
    return contracts


@router.delete("/{contract_id}")
@handle_errors
async def delete_contract(contract_id: str) -> Dict[str, Any]:
    """
    Delete a generated contract.
    Architettura MongoDB-first: elimina dal database.
    """
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    # Delete record from MongoDB (architettura MongoDB-first)
    await db["employee_contracts"].delete_one({"id": contract_id})
    
    # Cleanup opzionale: tenta eliminazione file locale se esiste (per retrocompatibilità)
    filepath = contract.get("filepath")
    if filepath:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except Exception:
            pass  # Ignora errori filesystem, il dato importante è su MongoDB
    
    return {"success": True, "message": "Contratto eliminato"}


# ---------------------------------------------------------------------------
# Invio del contratto/regolamento per email al dipendente + presa visione
# ---------------------------------------------------------------------------
def _smtp_send(to_addr: str, subject: str, body: str, allegati: List[Dict[str, Any]]) -> None:
    """Invio email con allegati via SMTP. Credenziali SOLO da env Render."""
    host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    port = int(os.getenv("SMTP_PORT", "465"))
    user = os.getenv("SMTP_EMAIL") or os.getenv("SMTP_USER")
    pwd = os.getenv("SMTP_PASSWORD")
    if not (user and pwd):
        raise HTTPException(503, "Email non configurata: imposta SMTP_EMAIL e SMTP_PASSWORD in env Render.")
    msg = EmailMessage()
    msg["From"] = user
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.set_content(body)
    for a in allegati:
        msg.add_attachment(a["data"], maintype="application",
                           subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                           filename=a["filename"])
    if port == 465:
        with smtplib.SMTP_SSL(host, port, context=ssl.create_default_context()) as s:
            s.login(user, pwd)
            s.send_message(msg)
    else:
        with smtplib.SMTP(host, port) as s:
            s.starttls(context=ssl.create_default_context())
            s.login(user, pwd)
            s.send_message(msg)


# Documenti accessori che accompagnano SEMPRE il contratto (da sottoscrivere):
# informativa 152/1997, informativa privacy e regolamento interno.
DOC_ACCESSORI = ["informativa_152", "informativa_privacy", "regolamento"]


async def _raccogli_documenti(db, contract: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Contratto + ultimi documenti accessori generati per il dipendente
    (informativa 152, privacy, regolamento), nell'ordine di sottoscrizione.

    Ritorna [{filename, data(bytes), tipo}]. Solo i documenti effettivamente
    generati per quel dipendente vengono inclusi.
    """
    docs = [{"filename": contract.get("filename", "contratto.docx"),
             "data": base64.b64decode(contract["file_data"]),
             "tipo": contract.get("contract_type", "contratto")}]
    emp_id = contract.get("employee_id")
    for tipo in DOC_ACCESSORI:
        d = await db["employee_contracts"].find_one(
            {"employee_id": emp_id, "contract_type": tipo},
            {"_id": 0, "file_data": 1, "filename": 1}, sort=[("generated_at", -1)])
        if d and d.get("file_data"):
            docs.append({"filename": d.get("filename", f"{tipo}.docx"),
                         "data": base64.b64decode(d["file_data"]), "tipo": tipo})
    return docs


@router.post("/send/{contract_id}")
@handle_errors
async def send_contract(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Invia al dipendente via email il contratto INSIEME a regolamento, privacy
    e informativa 152 (quelli generati per lui), da sottoscrivere."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    emp = await db[Collections.EMPLOYEES].find_one(
        {"id": contract.get("employee_id")},
        {"_id": 0, "nome": 1, "cognome": 1, "email": 1}) or {}
    to_addr = (data.get("email") or emp.get("email") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Email del dipendente mancante: inseriscila in anagrafica.")
    if not contract.get("file_data"):
        raise HTTPException(404, "File del contratto non disponibile")

    documenti = await _raccogli_documenti(db, contract)
    allegati = [{"filename": d["filename"], "data": d["data"]} for d in documenti]
    mancanti = [t for t in DOC_ACCESSORI if t not in {d["tipo"] for d in documenti}]

    nome = f"{emp.get('nome','')} {emp.get('cognome','')}".strip() or "Gentile collaboratore"
    elenco = ", ".join(d["filename"] for d in documenti)
    corpo = (
        f"Gentile {nome},\n\n"
        f"in allegato trova i documenti da sottoscrivere per l'assunzione: {elenco}.\n"
        f"La preghiamo di prenderne visione e di firmarli per accettazione.\n\n"
        f"Ceraldi Group S.r.l."
    )
    import asyncio
    await asyncio.to_thread(_smtp_send, to_addr,
                            f"Documenti di assunzione — {contract.get('contract_name','')}", corpo, allegati)
    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {"inviato_il": datetime.now(timezone.utc).isoformat(), "inviato_a": to_addr}})
    return {"ok": True, "inviato_a": to_addr, "allegati": len(allegati),
            "documenti": [d["filename"] for d in documenti], "accessori_mancanti": mancanti}


# ---------------------------------------------------------------------------
# Firma digitale via OpenAPI: marca temporale -> eSignature (FES+OTP) -> PEC
# Stato nel fascicolo: inviato -> firmato -> accettato.
# ---------------------------------------------------------------------------
def _docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """Converte un .docx in PDF tramite LibreOffice headless.

    Render non include LibreOffice di default: se `soffice`/`libreoffice` non è
    disponibile viene sollevato un errore chiaro e azionabile.
    """
    import subprocess
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise HTTPException(
            501,
            "Conversione PDF non disponibile: LibreOffice (soffice) non è "
            "installato sull'ambiente. Aggiungilo al deploy oppure usa un "
            "servizio di conversione esterno.")
    workdir = tempfile.mkdtemp(prefix="docx2pdf_")
    src = os.path.join(workdir, "contratto.docx")
    with open(src, "wb") as f:
        f.write(docx_bytes)
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", workdir, src],
            check=True, capture_output=True, timeout=120)
        pdf_path = os.path.join(workdir, "contratto.pdf")
        if not os.path.exists(pdf_path):
            raise HTTPException(500, "Conversione PDF fallita (output assente).")
        with open(pdf_path, "rb") as f:
            return f.read()
    except subprocess.CalledProcessError as e:
        raise HTTPException(500, f"Conversione PDF fallita: {e.stderr.decode('utf-8', 'ignore')[:300]}")
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


async def _get_contract_pdf(contract: Dict[str, Any]) -> bytes:
    """Ritorna il PDF del contratto: usa quello già marcato/firmato se presente,
    altrimenti converte il .docx generato."""
    if contract.get("pdf_data"):
        return base64.b64decode(contract["pdf_data"])
    if not contract.get("file_data"):
        raise HTTPException(404, "File del contratto non disponibile")
    return _docx_bytes_to_pdf(base64.b64decode(contract["file_data"]))


def _bundle_to_pdf(documenti: List[Dict[str, Any]]) -> bytes:
    """Converte ogni .docx (contratto + accessori) in PDF e li unisce in un unico
    PDF: così la firma per accettazione copre tutti i documenti in una sola volta."""
    from PyPDF2 import PdfMerger
    merger = PdfMerger()
    for d in documenti:
        merger.append(io.BytesIO(_docx_bytes_to_pdf(d["data"])))
    out = io.BytesIO()
    merger.write(out)
    merger.close()
    return out.getvalue()


@router.post("/sign/{contract_id}")
@handle_errors
async def avvia_firma(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Avvia il flusso di firma: converte in PDF, applica marca temporale e crea
    la richiesta di eSignature (FES con OTP) verso il dipendente."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    emp = await db[Collections.EMPLOYEES].find_one(
        {"id": contract.get("employee_id")},
        {"_id": 0, "nome": 1, "cognome": 1, "email": 1, "telefono": 1, "cellulare": 1}) or {}
    to_addr = (data.get("email") or emp.get("email") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Email del dipendente mancante: inseriscila in anagrafica.")
    phone = (data.get("phone") or emp.get("cellulare") or emp.get("telefono") or "").strip()
    nome = f"{emp.get('nome','')} {emp.get('cognome','')}".strip() or "Dipendente"

    client = get_client()
    if not client.configured:
        raise HTTPException(503, "OpenAPI non configurato: imposta OPENAPI_CLIENT_ID e "
                                 "OPENAPI_CLIENT_SECRET nelle env di Render.")
    import asyncio
    documenti = await _raccogli_documenti(db, contract)
    fname = (contract.get("filename", "documenti").rsplit(".", 1)[0]) + "_assunzione.pdf"
    try:
        # Unico PDF con contratto + regolamento + privacy + informativa 152.
        pdf_bytes = await asyncio.to_thread(_bundle_to_pdf, documenti)
        # 1) Marca temporale (data certa)
        ts = await client.apply_timestamp(pdf_bytes, filename=fname)
        # 2) eSignature FES con OTP verso il dipendente (firma unica su tutti i documenti)
        sig = await client.create_signature_request(
            pdf_bytes, signer_name=nome, signer_email=to_addr, signer_phone=phone,
            title=f"Documenti di assunzione — {contract.get('contract_name','')}",
            filename=fname)
    except OpenAPIConfigError as e:
        raise HTTPException(503, str(e))
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    req_id = sig.get("id") or sig.get("request_id")
    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "pdf_data": base64.b64encode(pdf_bytes).decode("ascii"),
            "firma_stato": "inviato",
            "firma_request_id": req_id,
            "marca_temporale": ts,
            "firma_inviata_a": to_addr,
            "firma_inviata_il": datetime.now(timezone.utc).isoformat(),
            "firma_documenti": [d["filename"] for d in documenti],
        }})
    return {"ok": True, "stato": "inviato", "request_id": req_id, "firmatario": to_addr,
            "documenti": [d["filename"] for d in documenti]}


@router.get("/sign/{contract_id}/status")
@handle_errors
async def stato_firma(contract_id: str) -> Dict[str, Any]:
    """Interroga OpenAPI sullo stato della firma e aggiorna il fascicolo.
    Se firmato, salva il PDF firmato e passa lo stato a 'firmato'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    req_id = contract.get("firma_request_id")
    if not req_id:
        return {"ok": True, "stato": contract.get("firma_stato") or "non_avviato"}

    client = get_client()
    try:
        res = await client.get_signature_status(req_id)
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    stato_raw = str(res.get("status") or "").lower()
    updates: Dict[str, Any] = {"firma_check_il": datetime.now(timezone.utc).isoformat()}
    firmato = stato_raw in ("signed", "completed", "firmato")
    if firmato:
        updates["firma_stato"] = "firmato"
        signed = res.get("signed_document") or res.get("signed_pdf") or {}
        content = signed.get("content") if isinstance(signed, dict) else signed
        if content:
            updates["pdf_data"] = content  # base64 del PDF firmato
            updates["firmato_il"] = datetime.now(timezone.utc).isoformat()
    await db["employee_contracts"].update_one({"id": contract_id}, {"$set": updates})
    return {"ok": True, "stato": updates.get("firma_stato", contract.get("firma_stato")), "provider": stato_raw}


@router.post("/pec/{contract_id}")
@handle_errors
async def invia_pec(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Invia via PEC il contratto (firmato se disponibile) con ricevuta a data certa.
    Porta lo stato a 'accettato'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    to_addr = (data.get("pec") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Indirizzo PEC destinatario mancante.")
    pdf_bytes = await _get_contract_pdf(contract)
    pdf_name = (contract.get("filename", "contratto.docx").rsplit(".", 1)[0]) + ".pdf"

    client = get_client()
    if not client.configured:
        raise HTTPException(503, "OpenAPI non configurato (OPENAPI_CLIENT_ID/SECRET in env Render).")
    try:
        res = await client.send_pec(
            to_addr=to_addr,
            subject=f"Contratto di assunzione — {contract.get('contract_name','')}",
            body="In allegato il contratto di assunzione con marca temporale e firma per accettazione.",
            attachments=[{"filename": pdf_name, "content": pdf_bytes}])
    except OpenAPIConfigError as e:
        raise HTTPException(503, str(e))
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "firma_stato": "accettato",
            "pec_inviata_a": to_addr,
            "pec_inviata_il": datetime.now(timezone.utc).isoformat(),
            "pec_messaggio": res,
        }})
    return {"ok": True, "stato": "accettato", "pec": to_addr}
